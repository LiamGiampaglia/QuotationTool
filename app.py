from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import streamlit as st
from docx import Document
import tempfile
from datetime import datetime
import openpyxl
import io


material_code_options = [
    # POWER CONSULTANCY
    "FSC11 EcoConsult Essential / MPS Walkthrough",
    "FSC719 Eco Audit Advanced Plus / MPS Enterprise",
    "FSC9002 Electrical Digital Twin",
    "FSC830 Re-Assessment Audit",
    "FSC612 EcoConsult Advanced",
    "FSC9007 Arc Flash_ETAP",
    "FSC9006 Protection Study_ETAP",
    "CON594 Fault Level Study",
    "CON607 Protection System Desig",
    "CON602 Earthing Design",
    "FSC5099 Relay Programming",
    "FSC5099 Pressure Rise Studies",
    "FSC5099 Network Design",
    "FSC5099 Substation Design",
    "FSC5099 Equipment Design",
    "SRV_ETAP_PQ_MOD Harmonic Surveys / Temp Monitoring",
    "FSC9011 ETAP Subscription",
    "SRV_PM_AUDIT_ADV EcoConsult Audit Power Monitoring audit",
    "SRV_PQ_AUDIT_ESS EcoConsult Audit PQ Essential",

    # PROCESS ELECTRIFICATION
    "PEPCONSULTING_0001 Process Electrification Project",

    # INDUSTRY & DT
    "GCRLIFECYCLECONSUL Industry Consulting (IDIBS)",
    "GCRDIGITRANSPRJCT Digital Transformation Project",

    # DATA CENTRE & COOLING
    "WCONSULTADV EcoConsult for Data Centers",

    # BMS
    "CON596 BMS Consultancy",

    # EV
    "EVS1AG eMobility Consultancy Audit",

    # ENERGY
    "CON200 Sust.Serv Smart Grid Elect.Audi",
    "CON201 Sust.Serv Smart Grid Elect.Design",
    "CON202 Sust.Serv EV Infrastructure Audit",
    "CON203 Sust.Serv Energy Efficiency Audit",
    "CON204 Sust.Serv Remote Energy Assessment",
    "CON205 Sust.Serv Metering Study",
    "CON206 Sust.Serv Energy Monitoring",

    # MICROGRID
    "SRVINAMGFEAS Microgrid Feasibility Study",
    "SRVINAMGDES Microgrid Design",

    # EXTRA
    "CON601 DON’T USE Consultancy Ancillary Offers",
    "CON595 Electrical Network Design"
]

def extract_rates(uploaded_file):

    wb = openpyxl.load_workbook(uploaded_file, data_only=True)
    ws = wb["PRICING SHEET"]

    
    rates = {
        "office_cost": ws["E97"].value or 0,
        "site_cost": ws["E98"].value or 0,
    
        "office_day": ws["E101"].value or 0,
        "office_evening": ws["F101"].value or 0,
        "office_weekend": ws["G101"].value or 0,
    
        "site_day": ws["E102"].value or 0,
        "site_evening": ws["F102"].value or 0,
        "site_weekend": ws["G102"].value or 0,
    
        "outside_m25": ws["E106"].value or 0,
        "inside_m25": ws["E107"].value or 0,
        "mileage": ws["E108"].value or 0
    }


    return rates



def generate_pricing_excel(uploaded_file):

    uploaded_file.seek(0)
    file_bytes = uploaded_file.read()
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=False)

    ws = wb["PRICING SHEET"]

    # ✅ Header info
    ws["C5"] = sap_description
    ws["C8"] = consultant_name
    ws["J5"] = currency
    # ✅ SAP Sheet
    sap_ws = wb["SAP INFO FORM"]

    sap_ws["D7"] = customer_contact_name
    sap_ws["D8"] = contact_tel
    sap_ws["D9"] = contact_email

    sap_ws["D17"] = bfo_opp_no
    sap_ws["D18"] = material_code_count

    sap_ws["D19"] = material_code_1
    sap_ws["D20"] = material_code_2
    sap_ws["D21"] = material_code_3
    
    sap_ws["E19"] = material_price_1
    sap_ws["E20"] = material_price_2 if material_code_count >= 2 else ""
    sap_ws["E21"] = material_price_3 if material_code_count == 3 else ""
    
    # ======================
    # ✅ LABOUR ROWS
    # ======================
    for i, lr in enumerate(st.session_state.labour_rows):
        if i >= 10:
            break
        row = 15 + i
        ws[f"C{row}"] = lr["description"]
        ws[f"D{row}"] = lr["office_day"]
        ws[f"E{row}"] = lr["site_day"]
        ws[f"F{row}"] = lr["office_evening"]
        ws[f"G{row}"] = lr["site_evening"]
        ws[f"H{row}"] = lr["office_weekend"]
        ws[f"I{row}"] = lr["site_weekend"]
    
    # ======================
    # ✅ OTHER COSTS
    # ======================
    for i, oc in enumerate(st.session_state.other_cost_rows):
        if i >= 6:
            break
        row = 39 + i
        ws[f"C{row}"] = oc["description"]
        ws[f"D{row}"] = oc["cost"]
        ws[f"E{row}"] = "GBP"
        ws[f"G{row}"] = oc["margin"] / 100
    
    # ======================
    # ✅ EXPENSES
    # ======================
    ws["D30"] = st.session_state.get("overnight_outside", 0)
    ws["D31"] = st.session_state.get("overnight_inside", 0)
    ws["D32"] = st.session_state.get("miles", 0)
    ws["D33"] = st.session_state.get("flights_cost", 0)
    
    # ======================
    # ✅ DISCOUNT
    # ======================
    ws["D62"] = st.session_state.get("discount_pct", 0) / 100
    
    return wb



def load_excel_into_session(uploaded_file):

    def safe_num(val):
        if val is None:
            return 0
        try:
            return float(val)
        except:
            try:
                return float(str(val).strip())
            except:
                return 0

    wb = openpyxl.load_workbook(uploaded_file, data_only=True)
    ws = wb["PRICING SHEET"]

    # ==========================
    # ✅ LABOUR ROWS
    # ==========================
    labour_rows = []

    for row in range(15, 25):
        desc = ws[f"C{row}"].value
        if desc:
            labour_rows.append({
                "description": str(desc),
                "office_day": safe_num(ws[f"D{row}"].value),
                "site_day": safe_num(ws[f"E{row}"].value),
                "office_evening": safe_num(ws[f"F{row}"].value),
                "site_evening": safe_num(ws[f"G{row}"].value),
                "office_weekend": safe_num(ws[f"H{row}"].value),
                "site_weekend": safe_num(ws[f"I{row}"].value),
            })

    # ==========================
    # ✅ OTHER COSTS
    # ==========================
    other_cost_rows = []

    for row in range(39, 45):
        desc = ws[f"C{row}"].value
        if desc:
            other_cost_rows.append({
                "description": str(desc),
                "cost": safe_num(ws[f"D{row}"].value),
                "margin": safe_num(ws[f"G{row}"].value) * 100,
                "selling": 0.0
            })

    # ==========================
    # ✅ EXPENSES
    # ==========================
    expenses = {
        "overnight_outside": safe_num(ws["D30"].value),
        "overnight_inside": safe_num(ws["D31"].value),
        "miles": safe_num(ws["D32"].value),
        "flights_cost": safe_num(ws["D33"].value),
    }

    discount_raw = safe_num(ws["D62"].value)
    discount_pct = discount_raw * 100

    return labour_rows, other_cost_rows, expenses, discount_pct

def load_billing_milestones(uploaded_file):

    def safe_num(val):
        try:
            return float(val)
        except:
            return 0.0

    def safe_date(val):
        if isinstance(val, datetime):
            return val
        return datetime.today()

    wb = openpyxl.load_workbook(uploaded_file, data_only=True)
    sap_ws = wb["SAP INFO FORM"]

    milestone_count = int(sap_ws["D37"].value or 0)

    billing_values = []
    billing_dates = []

    for i in range(milestone_count):

        value_cell = f"D{38 + i * 2}"
        date_cell = f"D{39 + i * 2}"

        value = safe_num(sap_ws[value_cell].value)
        date = safe_date(sap_ws[date_cell].value)

        billing_values.append(value)
        billing_dates.append(date)

    return milestone_count, billing_values, billing_dates
    

# ==========================
# SESSION STATE
# ==========================


if "works_list" not in st.session_state:
    st.session_state.works_list = []

if "other_cost_rows" not in st.session_state:
    st.session_state.other_cost_rows = []

if "labour_rows" not in st.session_state:
    st.session_state.labour_rows = []

if "price1" not in st.session_state:
    st.session_state.price1 = 0.0

if st.session_state.get("do_autofill", False):

    if "total_price" in st.session_state:
        st.session_state.price1 = float(st.session_state.total_price)
        st.session_state.bm_value_0 = float(st.session_state.total_price)

    st.session_state.do_autofill = False


if "currency" not in st.session_state:
    st.session_state.currency = "GBP"

currency = st.session_state.currency

if currency == "EUR":
    fx_rate = 1.16
    currency_symbol = "€"
else:
    fx_rate = 1
    currency_symbol = "£"

rates = {
    "office_day": 0,
    "site_day": 0,
    "office_evening": 0,
    "site_evening": 0,
    "office_weekend": 0,
    "site_weekend": 0
}

expenses_total = 0
discount_factor = 1


# ==========================
# PAGE SETUP
# ==========================
st.set_page_config(page_title="Energy Quote Tool", layout="centered")
st.title("Consultancy Quote Generator")

st.markdown("---")
# ==========================
# 📊 COST SHEET UPLOAD
# ==========================
st.subheader("📊 Cost Sheet")

with st.expander("💡 How to fill in", expanded=False):
    st.markdown("""
    ### How to fill in

    This section determines how the application handles pricing and how much automation is applied when generating the quotation and pricing sheet.

    Please select one of the following Template Modes:

    **1. No Pricing Template Uploaded**  
    - Use this option if you only require a quotation document  
    - All pricing and work items must be entered manually within the app  
    - No Excel pricing sheet will be used  

    **2. Blank Pricing Template Uploaded**  
    - Upload a blank pricing template from the following location:  
      https://schneiderelectric.sharepoint.com/sites/ConsultancyQdriveinternalGroup/Shared%20Documents/Forms/AllItems.aspx?csf=1&web=1&e=zUgqAU&CID=1df172e3%2Df1e1%2D40c7%2D99cd%2D639cc9bcd032&FolderCTID=0x012000EB746CE09F8B034EA74F90EDFEBE8CFD&id=%2Fsites%2FConsultancyQdriveinternalGroup%2FShared%20Documents%2FGeneral%2F03%20QMS%20Documents%2F04%20Forms%2FForQ%5FUKI%5FCNS01%20Pricing%20Template  
    - Complete the sections within the tool to generate:
        - A quotation document  
        - A fully populated pricing Excel file  
    - This removes the need to manually complete the Excel separately  

    **3. Pre-Populated Template Uploaded**  
    - Upload a completed pricing template  
    - All pricing data will be automatically pulled into the tool  
    - You only need to complete the quotation details  
    - The system will:
        - Populate calculations automatically  
        - Generate the quotation document  
        - Update the pricing Excel  

    ⚠️ Selecting the correct mode is essential for ensuring pricing and outputs behave as expected.
    """)
    
template_mode = st.selectbox(
    "Select Template Mode",
    [
        "No Pricing Template Uploaded",
        "Blank Pricing Template Uploaded",
        "Pre-Populated Template Uploaded"
    ]
)

uploaded_file = st.file_uploader(
    "Upload Pricing Template",
    type=["xlsx"]
)
if (
    template_mode == "Pre-Populated Template Uploaded"
    and uploaded_file is not None
    and not st.session_state.get("excel_loaded", False)
):

    
    for i in range(10):
        st.session_state.pop(f"bm_value_{i}", None)
        st.session_state.pop(f"bm_date_{i}", None)
    
    # ✅ RESET SELECTBOX
    st.session_state.pop("billing_count_selectbox", None)
    
    # ✅ LOAD FROM EXCEL
    bm_count, bm_values, bm_dates = load_billing_milestones(uploaded_file)
    
    # ✅ STORE COUNT (THIS IS KEY)
    st.session_state.billing_milestone_count_loaded = bm_count


    for i in range(bm_count):
        st.session_state[f"bm_value_{i}"] = bm_values[i]
        st.session_state[f"bm_date_{i}"] = bm_dates[i]

    # ✅ LOAD MAIN EXCEL DATA
    labour_rows, other_cost_rows, expenses, discount_pct = load_excel_into_session(uploaded_file)

    st.session_state.labour_rows = labour_rows
    st.session_state.other_cost_rows = other_cost_rows

    st.session_state.overnight_outside = expenses["overnight_outside"]
    st.session_state.overnight_inside = expenses["overnight_inside"]
    st.session_state.miles = expenses["miles"]
    st.session_state.flights_cost = expenses["flights_cost"]
    st.session_state.discount_pct = discount_pct
    st.session_state.excel_loaded = True
    st.success("✅ Excel data including billing milestones loaded into calculator")
    st.rerun()


st.markdown("---")

# ==========================
# DISCIPLINE SELECTION
# ==========================
with st.expander("Quotation Type", expanded=False):
    with st.expander("💡 How to fill in", expanded=False):
        st.markdown("""
        **This section determines the quotation template that will be used to generate the document.**
    
        Please follow the steps below:
    
        - Select a **Discipline** from the available options  
        - Then select the **Quote Type** document you would like to generate  
    
        The selected quote type will determine:
        - The structure of the quotation document  
        - The template used when generating the final Word output  
    
        ✅ Ensure the correct discipline and quote type are selected before proceeding.
    
        ⚠️ Selecting the wrong option may result in the incorrect template being used.
        """)

    discipline_options = ["Energy", "Power", "Microgrid", "Data Centre"]
    discipline = st.selectbox("Select Discipline", discipline_options)
    
    
    # ==========================
    # QUOTE TYPE BY DISCIPLINE
    # ==========================
    
    quote_options_by_discipline = {
    
        "Energy": {
            "Energy Efficiency Audit": "templates/Energy Efficiency Audit Template.docx",
            "Metering Assessment": "templates/Metering Assessment Template.docx",
            "EE and Metering": "templates/EE Audit and Metering Template.docx",
            "ESOS P4": "templates/ESOS P4 Template.docx",
            "ESOS P4 and Transport": "templates/ESOS P4 and Transport Template.docx",
            "ESOS P4 Transport": "templates/ESOS P4 Transport Template.docx",
            "Energy Performance Certificate": "templates/Energy Performance Certificate Template.docx"
        },
    
        "Power": {},       # ✅ ready for future
        "Microgrid": {},   # ✅ ready for future
        "Data Centre": {}  # ✅ ready for future
    }
    
    
    quote_options = quote_options_by_discipline.get(discipline, {})
    
    if quote_options:
        quote_type = st.selectbox("Select Quote Type", list(quote_options.keys()))
        st.session_state.selected_quote_type = quote_type
    else:
        st.warning("No quote types available for this discipline yet.")
        quote_type = None
    
st.markdown("---")

# ==========================
# INPUT FIELDS
# ==========================

with st.expander("📄 Template Info", expanded=False):
    with st.expander("💡 How to fill in", expanded=False):
        st.markdown("""
        **All fields in this section must be completed to ensure the quote template is populated correctly.**
    
        Please provide accurate project and customer information, including:
        - Company name
        - Project name (EN Number)  
        - Number and location of sites (Location of sites to be filled in like the example)  
        - Consultant Name (Main Consultant)
        - Number of Consultants 
    
        ⚠️ Missing or incomplete information may result in sections of the generated quote not being filled correctly.
        """)
    col1, col2 = st.columns(2)
    
    with col1:
        customer_name = st.text_input("Company Name")
        number_of_sites = st.text_input("Number of Sites")
    
        site_name = st.text_input(
            "Site Name",
            placeholder="e.g. Coventry, London and Warrington"
        )
        st.caption("If multiple sites, use this format: Coventry, London and Warrington")
    
    with col2:
        project_name = st.text_input("Project Name")
    
        consultant_name = st.text_input(
            "Consultant Name",
        )
        number_of_consultants = st.number_input(
            "Number of Consultants",
            min_value=1,
            max_value=50,
            value=1
        )

# ==========================
# 🏢 CUSTOMER DETAILS
# ==========================
st.markdown("---")
with st.expander("🏢 Customer Details", expanded=False):
    
    with st.expander("💡 How to fill in", expanded=False):
        st.markdown("""
        **All address fields in this section must be completed to ensure the customer details are correctly populated in the quote.**
    
        Please provide:
        - Full address (Address Line 1, Address Line 2 if applicable, City, and Postcode)
    
        **Contact Name is optional** and can be left blank if unknown.  
        If left blank, the quote will default to a generic title.
    
        ⚠️ Missing address details may result in incomplete customer information in the generated document.
        """)


    col1, col2 = st.columns(2)
    
    with col1:
        address_line_1 = st.text_input("Address Line 1")
        address_line_2 = st.text_input("Address Line 2")
    
    with col2:
        city = st.text_input("City")
        postcode = st.text_input("Postcode")
    
    # Contact name (optional)
    contact_name = st.text_input("Contact Name (leave blank for default)")
st.markdown("---")
# ==========================
# 💰 LIVE COST CALCULATOR
# ==========================
if uploaded_file is not None:
    with st.expander("💰 Live Cost Calculator", expanded=True):
        with st.expander("💡 How to fill in", expanded=False):
            st.markdown("""
            ### How to fill in
        
            This section is used to calculate the total project cost and selling price based on labour, expenses, and other costs.
        
            The process depends on the Template Mode selected:
        
            ---
        
            **1. Pre-Populated Template Uploaded**
        
            - All labour rates, costs, and values are automatically pulled from the uploaded Excel file  
            - Your role is to:
                - ✅ Review the imported data carefully  
                - ✅ Check that all values match the original Excel file  
                - ✅ Adjust only if necessary  
        
            👉 No manual data entry is required unless corrections are needed  
        
            ---
        
            **2. Blank Pricing Template Uploaded**
        
            You must complete this section fully to build the pricing.
        
            Follow the steps below:
        
            **Step 1 – Select Currency**
            - Choose the required currency (**GBP** or **EUR**)  
            - All calculations and totals will update automatically  
        
            **Step 2 – Add Labour Items**
            - Click **Add Labour Row**  
            - Enter:
                - Description of the work  
                - Hours for each category (Office / Site / Evening / Weekend)  
            - Repeat for all work activities  
        
            **Step 3 – Add Other Costs**
            - Click **Add Cost Row**  
            - Enter:
                - Description  
                - Cost (£)  
                - Margin (%)  
            - The selling price will be calculated automatically  
        
            **Step 4 – Enter Expenses**
            - Input:
                - Overnight stays (Inside / Outside M25)  
                - Mileage  
                - Travel costs (Flights / Rail)  
        
            **Step 5 – Apply Discount (if required)**
            - Enter a discount percentage  
            - The final selling price will be adjusted automatically  
        
            **Step 6 – Auto Fill Pricing Fields (Optional)**
            - Click **Auto Fill Pricing Fields**  
            - This will:
                - Populate material pricing values automatically  
                - Align pricing with the calculated project total  
        
            ✅ This helps speed up completion of the Pricing Sheet Info section  
        
            ---
        
            **What the Calculator Does**
        
            The tool will automatically:
            - Apply labour rates from the pricing template  
            - Calculate selling prices and costs  
            - Apply peer review uplift  
            - Calculate margins and final project value  
            - Convert values if EUR is selected  
        
            ---
        
            ✅ The final total is used throughout the tool, including:
            - Works & Pricing  
            - Billing Milestones  
            - Payment Terms  
            - Final quotation document  
        
            ---
        
            ⚠️ Ensure all values are accurate, as errors here will impact pricing, margins, and customer-facing outputs.
            """)

        rates = extract_rates(uploaded_file)
        st.write("DEBUG RATES:", rates)
    
        st.markdown("---")
        st.subheader("💰 Live Cost Calculator")
    
        currency = st.selectbox(
            "Currency",
            ["GBP", "EUR"],
            key="currency"
        )
        
        currency = st.session_state.currency
        
        if currency == "EUR":
            fx_rate = 1.16
            currency_symbol = "€"
        else:
            fx_rate = 1
            currency_symbol = "£"
    
        # Inputs
        st.markdown("### Labour Hours (Detailed)")
    
    # ✅ Add row button
        if st.button("➕ Add Labour Row"):
            st.session_state.labour_rows.append({
                "description": "",
                "office_day": 0.0,
                "site_day": 0.0,
                "office_evening": 0.0,
                "site_evening": 0.0,
                "office_weekend": 0.0,
                "site_weekend": 0.0
            })
        
        # ✅ Display rows
        total_office_day = 0
        total_site_day = 0
        total_office_evening = 0
        total_site_evening = 0
        total_office_weekend = 0
        total_site_weekend = 0
        
        for i, row in enumerate(st.session_state.labour_rows):
        
            st.markdown(f"#### Work Item {i+1}")
        
            row["description"] = st.text_input(
                "Description",
                value=row["description"],
                key=f"payment_desc_{i}"
            )
        
            col1, col2, col3 = st.columns(3)
        
            
            with col1:
                

                row["office_day"] = st.number_input(
                    "Office Hours (Mon-Fri 0800 to 1700)",
                    value=float(row.get("office_day", 0.0)),
                    key=f"od_{i}"
                )


                
                row["site_day"] = st.number_input(
                    "On Site Hours (Mon-Fri 0800 to 1700)",
                    value=float(row.get("site_day", 0.0)),
                    key=f"sd_{i}"
                )

            
            with col2:
                
                row["office_evening"] = st.number_input(
                    "Office Hours (Mon-Fri 1700 to 2400)",
                    value=float(row.get("office_evening", 0.0)),
                    key=f"oe_{i}"
                )

               
                row["site_evening"] = st.number_input(
                    "On Site Hours (Mon-Fri 1700 to 2400)",
                    value=float(row.get("site_evening", 0.0)),
                    key=f"se_{i}"
                )

            
            with col3:
                
                row["office_weekend"] = st.number_input(
                    "Office Hours (Sat&Sun 0800 to 2400)",
                    value=float(row.get("office_weekend", 0.0)),
                    key=f"ow_{i}"
                )

                
                row["site_weekend"] = st.number_input(
                    "On Site Hours (Sat & Sun 0800 to 2400)",
                    value=float(row.get("site_weekend", 0.0)),
                    key=f"sw_{i}"
                )

    
        
            # ✅ Accumulate totals
            total_office_day += row["office_day"]
            total_site_day += row["site_day"]
            total_office_evening += row["office_evening"]
            total_site_evening += row["site_evening"]
            total_office_weekend += row["office_weekend"]
            total_site_weekend += row["site_weekend"]
        
            st.markdown("---")
       
        st.markdown("### Other Costs (Detailed)")
    
        # ✅ Add row button
        if st.button("➕ Add Cost Row"):
            st.session_state.other_cost_rows.append({
                "description": "",
                "cost": 0.0,
                "margin": 0.0,
                "selling": 0.0
            })
        
        
        total_other_cost = 0
        total_other_selling = 0
        
        # ✅ Display rows
        for i, row in enumerate(st.session_state.other_cost_rows):
        
            st.markdown(f"#### Cost Item {i+1}")
        
            row["description"] = st.text_input(
                "Description",
                value=row["description"],
                key=f"other_desc_{i}"
            )
        
            col1, col2 = st.columns(2)
        
            with col1: 
                row["cost"] = st.number_input(
                    "Cost (£)",
                    value=float(row.get("cost", 0.0)),
                    key=f"other_cost_{i}"
                )

            with col2:
                row["margin"] = st.number_input(
                    "Margin (%)",
                    min_value=0.0,
                    max_value=100.0,
                    value=float(row.get("margin", 0.0)),
                    key=f"other_margin_{i}"
                )
        
            # ✅ Selling calculation
            if row["margin"] < 100:
                row["selling"] = row["cost"] / (1 - row["margin"] / 100)
            else:
                row["selling"] = 0
        
            st.write(f"Selling Price: £{row['selling']:,.2f}")
        
            total_other_cost += row["cost"]
            total_other_selling += row["selling"]
        
            st.markdown("---")
        
        # ✅ Totals (like Excel bottom row)
    
        peer_review_hours = 0.1 * total_office_day
        
        st.write("### Totals")
        st.write(f"Office Day: {total_office_day}")
        st.write(f"Peer Review (Office): {peer_review_hours}")
        st.write(f"Site Day: {total_site_day}")
        
        st.write(f"Office Evening: {total_office_evening}")
        st.write(f"Site Evening: {total_site_evening}")
        
        st.write(f"Office Weekend: {total_office_weekend}")
        st.write(f"Site Weekend: {total_site_weekend}")
        
        # ✅ SELLING LABOUR (NEW)
        labour_total = (
            (total_office_day * rates["office_day"]) +
            (total_site_day * rates["site_day"]) +
        
            (total_office_evening * rates["office_evening"]) +
            (total_site_evening * rates["site_evening"]) +
        
            (total_office_weekend * rates["office_weekend"]) +
            (total_site_weekend * rates["site_weekend"])
        )
        
        # ✅ Peer review
        peer_review = 0.1 * total_office_day * rates["office_day"]
        peer_review_factor = 0.1
        
        labour_total += peer_review
    
        st.markdown("### Expenses & Costs")
    
        col1, col2 = st.columns(2)
        
        with col1:
            
            overnight_outside = st.number_input(
                "Overnight Stays (Outside M25)",
                value=st.session_state.get("overnight_outside", 0)
            )
            
            overnight_inside = st.number_input(
                "Overnight Stays (Inside M25)",
                value=st.session_state.get("overnight_inside", 0)
            )
            
            miles = st.number_input(
                "Mileage (miles)",
                value=st.session_state.get("miles", 0)
            )
            
            flights_cost = st.number_input(
                "Flights / Rail (£)",
                value=st.session_state.get("flights_cost", 0.0)
            )
            
        st.session_state.overnight_outside = overnight_outside
        st.session_state.overnight_inside = overnight_inside
        st.session_state.miles = miles
        st.session_state.flights_cost = flights_cost

    
        # ✅ EXPENSES (SELLING)
        expenses_total = (
            overnight_outside * rates.get("outside_m25", 0)
            + overnight_inside * rates.get("inside_m25", 0)
            + miles * rates.get("mileage", 0)
            + flights_cost * 1.15
        )
 
        # ==========================
        # ✅ DISCOUNT SECTION
        # ==========================
        st.markdown("### Discount")
        
        
        discount_pct = st.number_input(
            "Discount (%)",
            min_value=0.0,
            max_value=100.0,
            value=float(st.session_state.get("discount_pct", 0.0))
        )
        st.session_state.discount_pct = discount_pct

        
        discount_factor = 1 - (discount_pct / 100)
           
            
        # ==========================
        # Other Costs
        # ==========================
            
        
        other_cost = total_other_cost
        other_cost_selling = total_other_selling
    
            
        # ==========================
        # Final Total
        # ==========================
    
            
        # ✅ COST CALCULATION (from Excel logic)
        
        labour_cost = (
            (total_office_day * rates["office_cost"]) +
            (total_site_day * rates["site_cost"]) +
        
            (total_office_evening * rates["office_cost"]) +
            (total_site_evening * rates["site_cost"]) +
        
            (total_office_weekend * rates["office_cost"]) +
            (total_site_weekend * rates["site_cost"])
        )
        
        peer_review_cost = 0.1 * total_office_day * rates["office_cost"]
        
        labour_cost += peer_review_cost
    
    
        # ✅ Expense COST (not selling)
        expenses_cost = (
            overnight_outside * (rates.get("outside_m25", 0) / 1.15 if rates.get("outside_m25", 0) else 0)
            + overnight_inside * (rates.get("inside_m25", 0) / 1.15 if rates.get("inside_m25", 0) else 0)
            + miles * (rates.get("mileage", 0) / 1.675 if rates.get("mileage", 0) else 0)
            + flights_cost
        )
    
        total_cost = labour_cost + expenses_cost + other_cost
    
            
        subtotal = labour_total + expenses_total + other_cost_selling
            
        discount_value = subtotal * (discount_pct / 100)
            
        total_price = subtotal - discount_value
        st.session_state["final_total_price"] = total_price
        st.session_state.total_price = total_price    
            
        labour_total_fx = labour_total * fx_rate
        expenses_total_fx = expenses_total * fx_rate
        other_cost_selling_fx = other_cost_selling * fx_rate
        
        total_cost_fx = total_cost * fx_rate
        subtotal_fx = subtotal * fx_rate
        total_price_fx = total_price * fx_rate
        labour_cost_fx = labour_cost * fx_rate
        expenses_cost_fx = expenses_cost * fx_rate
        other_cost_fx = other_cost * fx_rate
    
        labour_selling_discounted = labour_total * discount_factor
        expenses_selling_discounted = expenses_total * discount_factor
        other_selling_discounted = other_cost_selling * discount_factor
    
        labour_selling_discounted_fx = labour_selling_discounted * fx_rate
        expenses_selling_discounted_fx = expenses_selling_discounted * fx_rate
        other_selling_discounted_fx = other_selling_discounted * fx_rate
    
        if subtotal > 0:
            margin_pct = (subtotal - total_cost) / subtotal * 100
        else:
            margin_pct = 0
    
        if total_price > 0:
            actual_margin_pct = (total_price - total_cost) / total_price * 100
        else:
            actual_margin_pct = 0
    
        
        st.markdown("### Breakdown")
    
        st.markdown("#### Labour")
        st.write(f"Selling: {currency_symbol}{labour_total_fx:,.2f}")
        st.write(f"Cost: {currency_symbol}{labour_cost_fx:,.2f}")
        st.write(f"Selling After Discount: {currency_symbol}{labour_selling_discounted_fx:,.2f}")
        
        st.markdown("---")
        
        st.markdown("#### Expenses")
        st.write(f"Selling: {currency_symbol}{expenses_total_fx:,.2f}")
        st.write(f"Cost: {currency_symbol}{expenses_cost_fx:,.2f}")
        st.write(f"Selling After Discount: {currency_symbol}{expenses_selling_discounted_fx:,.2f}")
        
        st.markdown("---")
        
        st.markdown("#### Other Costs")
        st.write(f"Selling: {currency_symbol}{other_cost_selling_fx:,.2f}")
        st.write(f"Cost: {currency_symbol}{other_cost_fx:,.2f}")
        st.write(f"Selling After Discount: {currency_symbol}{other_selling_discounted_fx:,.2f}")
        
        st.markdown("---")
          
    
        st.write(f"Total Cost: {currency_symbol}{total_cost_fx:,.2f}")
        st.write(f"Selling Price: {currency_symbol}{subtotal_fx:,.2f}")
    
        st.write(f"Margin (%): {margin_pct:.2f}%")
        
        st.write(f"Discount (%): {discount_pct:.2f}%")
        st.write(f"Actual Selling Price: {currency_symbol}{total_price_fx:,.2f}")
        st.write(f"Actual Margin (%): {actual_margin_pct:.2f}%")
        
        st.markdown("---")
        st.metric("Total Price", f"{currency_symbol}{total_price_fx:,.2f}")
    
        
        if st.button("Auto Fill Pricing Fields"):
            st.session_state["do_autofill"] = True

# ==========================
# WORKS INPUT
# ==========================
with st.expander("🛠️ Works & Pricing", expanded=False):
    if st.button("🔄 Auto Split 20% / 40% / 40%"):
    
        total = st.session_state.total_price
    
        # ✅ FORCE RESET OF SELECTBOX VALUE
        st.session_state.pop("billing_count_selectbox", None)
        st.session_state.pop("billing_milestone_count_saved", None)
    
        # ✅ SET OVERRIDE (THIS WILL NOW WORK)
        st.session_state.billing_milestone_count_override = 3
    
        # ✅ SET SPLIT VALUES
        st.session_state["split_values"] = [
            total * 0.20,
            total * 0.40,
            total * 0.40
        ]
    
        st.session_state.payment_terms_locked = False
    
        st.rerun()
    
    with st.expander("💡 How to fill in", expanded=False):
        st.markdown("""
        ### How to fill in
    
        This section defines the work items that will appear in the quotation document sent to the customer.
    
        The behaviour of this section depends on the Template Mode selected:
    
        **1. No Pricing Template Uploaded**  
        - Click **Add Work Item** to create a new line  
        - Ensure **Pricing Mode = Manual**  
        - Enter:
            - A **Description** (this appears in the quotation document)  
            - A **Price**  
        - To add multiple items, click **Add Work Item** again and repeat  
    
        ✅ Use this mode when creating a quotation manually without an Excel pricing file  
    
        ---
    
        **2. Blank Pricing Template Uploaded**  
        - Complete the **Live Cost Calculator** section first  
        - Labour, Other Costs, and Expenses will automatically populate into this section  
        - Leave items as **Auto** if you are happy with the structure  
    
        ✅ These items will appear automatically in the quotation document  
    
        - To customise:
            - Change **Pricing Mode → Manual**  
            - Edit the description or price  
            - Delete items if required  
    
        ✅ Use this option to control how work items are displayed in the quotation  
    
        ---
    
        **3. Pre-Populated Template Uploaded**  
        - This behaves the same as the Blank Pricing Template option  
        - All pricing and work items are automatically populated from the uploaded Excel file  
    
        ✅ You can still switch to Manual mode to customise descriptions and pricing if needed  
    
        ---
    
        ⚠️ Ensure descriptions are clear and accurate, as they will be presented directly to the customer in the final quotation.
        """)
    
    if st.button("➕ Add Work Item"):
        st.session_state.works_list.append({
            "description": "",
            "mode": "Manual",
            "manual_price": 0.0,
            "price": 0.0
        })
        
    combined_items = []
    
    # Labour
    for lr in st.session_state.labour_rows:
        if lr["description"]:
            combined_items.append({
                "description": lr["description"],
                "type": "labour",
                "data": lr
            })
    
    # Other Costs
    for oc in st.session_state.other_cost_rows:
        if oc["description"]:
            combined_items.append({
                "description": oc["description"],
                "type": "other",
                "data": oc
            })
    
    # Expenses
    if 'expenses_total' in locals() and expenses_total > 0:
        combined_items.append({
            "description": "Expenses",
            "type": "expenses"
        })
    
    
    # ✅ AUTO-POPULATE ONLY IF EMPTY OR SIZE CHANGED
    
    
    combined_items = []
    
    # Labour
    for lr in st.session_state.labour_rows:
        if lr["description"]:
            combined_items.append({
                "description": lr["description"],
                "type": "labour",
                "data": lr
            })
    
    # Other Costs
    for oc in st.session_state.other_cost_rows:
        if oc["description"]:
            combined_items.append({
                "description": oc["description"],
                "type": "other",
                "data": oc
            })
    
    # Expenses
    if 'expenses_total' in locals() and expenses_total > 0:
        combined_items.append({
            "description": "Expenses",
            "type": "expenses"
        })

# ✅ Populate works_list automatically (ONLY when empty OR size mismatch)

    if len(st.session_state.works_list) != len(combined_items):
        if uploaded_file is not None:
            st.session_state.works_list = []
    
            for item in combined_items:
                st.session_state.works_list.append({
                    "description": item["description"],
                    "mode": "Auto",
                    "manual_price": 0.0,
                    "price": 0.0
                })

    
    total_works_price = 0
    
    for i, work in enumerate(st.session_state.works_list):
    
        
        col1, col2 = st.columns([6, 1])
        
        with col1:
            st.markdown(f"#### Work Item {i+1}")
        
        with col2:
            if st.button("❌", key=f"delete_{i}"):
                st.session_state.works_list.pop(i)
                st.rerun()
    
            if uploaded_file is None:
                work["mode"] = "Manual"
                st.selectbox(
                    "Pricing Mode",
                    ["Manual"],
                    key=f"mode_{i}"
                )
            else:
                work["mode"] = st.selectbox(
                    "Pricing Mode",
                    ["Auto", "Manual"],
                    key=f"mode_{i}"
                )
    
        # ==========================
        # ✅ AUTO MODE
        # ==========================
        if work["mode"] == "Auto" and uploaded_file is not None:
    
            price = 0
    
            combined_items = []
    
            # ✅ Labour rows
            for lr in st.session_state.labour_rows:
                if lr["description"]:
                    combined_items.append({
                        "description": lr["description"],
                        "type": "labour",
                        "data": lr
                    })
    
            # ✅ Other costs
            for oc in st.session_state.other_cost_rows:
                if oc["description"]:
                    combined_items.append({
                        "description": oc["description"],
                        "type": "other",
                        "data": oc
                    })
    
            # ✅ Expenses
            if expenses_total > 0:
                combined_items.append({
                    "description": "Expenses",
                    "type": "expenses"
                })
    
            # ✅ Assign item based on position
            if i < len(combined_items):
                item = combined_items[i]
                work["description"] = item["description"]
            
                if item["type"] == "labour":
                    lr = item["data"]
            
                    price = (
                        # ✅ Normal labour
                        lr["office_day"] * rates["office_day"] +
                        lr["site_day"] * rates["site_day"] +
                        lr["office_evening"] * rates["office_evening"] +
                        lr["site_evening"] * rates["site_evening"] +
                        lr["office_weekend"] * rates["office_weekend"] +
                        lr["site_weekend"] * rates["site_weekend"]
            
                        # ✅ Peer review (correctly included ✅)
                        + (lr["office_day"] * rates["office_day"] * 0.1)
            
                    ) * discount_factor
    
                elif item["type"] == "other":
                    price = item["data"]["selling"] * discount_factor
    
                elif item["type"] == "expenses":
                    price = expenses_total * discount_factor
    
            else:
                work["description"] = ""
                price = 0
    
            # ✅ Locked description display
            st.text_input(
                "Description",
                value=work["description"],
                key=f"auto_desc_locked_{i}",
                disabled=True
            )
    
        # ==========================
        # ✅ MANUAL MODE
        # ==========================
        else:
            work["description"] = st.text_input(
                "Description",
                value=work["description"],
                key=f"work_desc_{i}"
            )
    
            work["manual_price"] = st.number_input(
                "Manual Price (£)",
                0.0,
                key=f"manual_{i}"
            )
    
            price = work["manual_price"]
    
        # ✅ FINAL PRICE DISPLAY (ONLY ONCE)
        work["price"] = price
        
        price_fx = price * fx_rate
        st.write(f"Price: {currency_symbol}{price_fx:,.2f}")
    
        total_works_price += price
    
        st.markdown("---")
    
    
    st.write("### Works Total")
    total_works_price_fx = total_works_price * fx_rate
    st.write(f"Total Works Price: {currency_symbol}{total_works_price_fx:,.2f}")

# ==========================
# 📋 PRICING SHEET INFO
# ==========================
if template_mode == "Blank Pricing Template Uploaded":

    st.markdown("---")
    with st.expander("📋 Pricing Sheet Info", expanded=False):
        with st.expander("💡 How to fill in", expanded=False):
            st.markdown("""
            ### How to fill in
        
            This section is used to populate the **Pricing Excel file** automatically.
        
            ⚠️ **Important:**  
            This section only needs to be completed when using the **Blank Pricing Template Uploaded** mode.
        
            ---
        
            **When to Use This Section**
        
            - ✅ Required for **Blank Pricing Template Uploaded**  
            - ❌ Not required for:
                - No Pricing Template Uploaded  
                - Pre-Populated Template Uploaded (data is already provided)  
        
            ---
        
            **What to Enter**
        
            - **SAP Description Name**  
              → A short description of the project for SAP  
        
            - **bFO Opportunity Number**  
              → The relevant opportunity reference  
        
            - **Customer Contact Details**
              - Name  
              - Telephone number  
              - Email address  
        
            - **Material Codes**
              - Select up to 3 material codes  
              - Enter the corresponding price for each  
        
            ---
        
            **How It Works**
        
            - The information entered here is written directly into the Pricing Excel file  
            - This removes the need to manually complete the SAP / pricing sheets  
            - The Excel file will be fully populated when downloaded  
        
            ---
        
            ✅ Ensure all required fields are completed accurately to avoid issues in SAP or pricing submission.  
        
            ⚠️ Incorrect or missing information may result in an incomplete or invalid pricing sheet.
            """)

        col1, col2 = st.columns(2)
    
        with col1:
            sap_description = st.text_input("SAP Description Name")  
            bfo_opp_no = st.text_input("bFO Opportunity No")
    
        # ✅ Keep these inside too (important)
        customer_contact_name = st.text_input("Customer Contact Name")
        contact_tel = st.text_input("Contact Tel No")
        contact_email = st.text_input("Contact Email")
    
        with col2:
    
            material_code_count = st.number_input(
                "No. of Material Codes",
                min_value=1,
                max_value=3,
                value=1
            )
    
            material_code_1 = st.selectbox(
                "Material Code No.1",
                material_code_options,
                key="mc1"
            )
    
            material_price_1 = st.number_input(
                "Price for Material Code 1 (£)",
                min_value=0.0,
                key="price1"
            )
    
            # ✅ Material Code 2
            material_code_2 = ""
            material_price_2 = 0.0
    
            if material_code_count >= 2:
                material_code_2 = st.selectbox(
                    "Material Code No.2",
                    material_code_options,
                    key="mc2"
                )
    
                material_price_2 = st.number_input(
                    "Price for Material Code 2 (£)",
                    min_value=0.0,
                    key="price2"
                )
    
            # ✅ Material Code 3
            material_code_3 = ""
            material_price_3 = 0.0
    
            if material_code_count == 3:
                material_code_3 = st.selectbox(
                    "Material Code No.3",
                    material_code_options,
                    key="mc3"
                )
    
                material_price_3 = st.number_input(
                    "Price for Material Code 3 (£)",
                    min_value=0.0,
                    key="price3"
                )


# ==========================
# 💰 BILLING MILESTONES (ONLY IF FILE UPLOADED)
# ==========================
if uploaded_file is not None:

    with st.expander("📆 Billing Milestones", expanded=False): 
        
        if st.button("🔄 Auto Split 20% / 40% / 40%"):
            
            total = st.session_state.total_price
        
            # ✅ Clear old selectbox value so override works
            st.session_state.pop("billing_count_selectbox", None)
            st.session_state.pop("billing_milestone_count_saved", None)
        
            # ✅ Set override
            st.session_state.billing_milestone_count_override = 3
        
            # ✅ Set split values
            st.session_state["split_values"] = [
                total * 0.20,
                total * 0.40,
                total * 0.40
            ]
        
            st.session_state.payment_terms_locked = False
        
            st.rerun()

        with st.expander("💡 How to fill in", expanded=False):
            st.markdown("""
            ### How to fill in
        
            This section defines how the project value will be split into billing stages.
        
            - Select the **number of billing milestones** required  
            - For each milestone, enter:
                - The **billing value (£)**  
                - The **planned billing date**
        
            The system will automatically:
            - Calculate each milestone as a percentage of the total project value  
            - Display how much each milestone contributes to the overall project  
        
            ---
        
            **General Guidance**
        
            - Ensure the total value across all milestones matches the total project value  
            - Milestones should reflect key project stages (e.g. mobilisation, mid-project, completion)  
            - Billing dates should align with the project timeline  
        
            ---
        
            **For Projects Above £50,000**
        
            - It is recommended to split billing into multiple stages  
            - A typical structure is:
                - 20% at mobilisation  
                - 40% mid-project  
                - 40% on completion  
        
            ✅ You can use the **Auto Split (20% / 40% / 40%)** button to apply this structure automatically  
        
            ---
        
            ⚠️ Incorrect milestone values may lead to inaccurate billing terms in the final quotation document.
            """)
        

        milestone_options = [1, 2, 3, 4, 5]
        
        
        default_count = st.session_state.get(
            "billing_milestone_count_override",
            st.session_state.get(
                "billing_milestone_count_saved",
                st.session_state.get("billing_milestone_count_loaded", 1)
            )
        )

        
        billing_milestone_count = st.selectbox(
            "No. of Billing Milestones",
            milestone_options,
            index=milestone_options.index(default_count),
            key="billing_count_selectbox"
        )
        st.session_state.billing_milestone_count_saved = billing_milestone_count
        
        
        if (
            "billing_milestone_count_override" in st.session_state
            and billing_milestone_count == st.session_state.billing_milestone_count_override
        ):
            del st.session_state.billing_milestone_count_override

        billing_values = []
        billing_dates = []

        total_project = st.session_state.get("total_price", 0)

        for i in range(billing_milestone_count):
        
            st.markdown(f"#### Billing Milestone {i+1}")
        
            key_name = f"bm_value_{i}"
        
            # ✅ FIXED INITIALISATION
            
            if "split_values" in st.session_state and i < len(st.session_state["split_values"]):
                st.session_state[key_name] = st.session_state["split_values"][i]
    
            else:
                if key_name not in st.session_state or st.session_state[key_name] == 0:
                    if total_project > 0:
                        st.session_state[key_name] = total_project / billing_milestone_count
                    else:
                        st.session_state[key_name] = 0.0

        
            value = st.number_input(
                f"Milestone {i+1} Value (£)",
                min_value=0.0,
                value=float(st.session_state[key_name]),
                key=key_name
            )
        
            date = st.date_input(
                f"Milestone {i+1} Planned Billing Date",
                value=st.session_state.get(f"bm_date_{i}", datetime.today()),
                key=f"bm_date_{i}"
            )
        
            
            if total_project > 0:
                percentage = (value / total_project) * 100
            
                # ✅ FIX: STORE PERCENTAGES PROPERLY
                if "billing_percentages" not in st.session_state:
                    st.session_state.billing_percentages = [0] * billing_milestone_count
            
                if len(st.session_state.billing_percentages) != billing_milestone_count:
                    st.session_state.billing_percentages = [0] * billing_milestone_count
            
                st.session_state.billing_percentages[i] = percentage  # ✅ IMPORTANT
            
                st.caption(f"📊 This milestone = {percentage:.0f}% of total project value")



if "total_price" in st.session_state:

    total_price_check = st.session_state.total_price

    if total_price_check > 50000:

        st.warning(
            f"⚠️ Project value is £{total_price_check:,.2f}. "
            "For projects above £50k, billing milestones should be split into multiple stages "
            "(e.g. 20% mobilisation, 40% mid-project, 40% completion)."
        )


# ==========================
# ✅ PAYMENT TERMS SYNC ENGINE
# ==========================

if "payment_terms" not in st.session_state:
    st.session_state.payment_terms = []

if "payment_terms_locked" not in st.session_state:
    st.session_state.payment_terms_locked = False

if "payment_override" not in st.session_state:
    st.session_state.payment_override = False


if "billing_percentages" in st.session_state:

    billing_pcts = st.session_state.billing_percentages

    # ✅ ONLY BUILD ONCE (DO NOT REBUILD EVERY RERUN)
    if not st.session_state.get("payment_terms_locked", False):

        if any(billing_pcts):
            st.session_state.payment_terms = [
                {
                    "percent": int(round(pct)),
                    "description": f"Milestone {i+1} completion"
                }
                for i, pct in enumerate(billing_pcts)
            ]
        else:
            st.session_state.payment_terms = [{
                "percent": 100,
                "description": "upon submittal of the report"
            }]

        # ✅ LOCK AFTER BUILD
        st.session_state.payment_terms_locked = True



    billing_pcts = st.session_state.billing_percentages


    # ✅ BUILD ONCE ONLY
    if all(pct == 0 for pct in billing_pcts):
        st.session_state.payment_terms = [{
            "percent": 100,
            "description": "upon submittal of the report"
        }]
    else:
        st.session_state.payment_terms = [
            {
                "percent": int(round(pct)),
                "description": f"Milestone {i+1} completion"
            }
            for i, pct in enumerate(billing_pcts)
        ]

# ==========================
# 💰 PAYMENT TERMS
# ==========================
st.markdown("---")


def update_payment_term(index):
    key = f"percent_{index}"
    if key in st.session_state:
        st.session_state.payment_terms[index]["percent"] = st.session_state[key]
        st.session_state.payment_override = True


with st.expander("💰 Payment Terms", expanded=False):
    with st.expander("💡 How to fill in", expanded=False):
        st.markdown("""
        ### How to fill in
    
        This section defines the payment terms that will appear in the quotation document sent to the customer.
    
        Each row represents a payment condition, combining:
        - A **percentage (%) of the total project value**  
        - A **description of when payment is due**
    
        ---
    
        **Blank Template Uploaded Auto-Generated Terms**
    
        - Payment terms are automatically created based on the **Billing Milestones** section for Blank Template Uploads  
        - Each milestone will populate a corresponding payment percentage  
        - Descriptions are generated based on milestone stages  
    
        ✅ This ensures payment terms align with your billing structure  
    
        ---
    
        **Manual Adjustments**
    
        - You can edit:
            - The **percentage (%)**  
            - The **description (payment condition)**  
        - Once edited, the system switches to **manual override mode**  
    
        ✅ This allows full control over how payment terms are presented  
    
        ---
    
        **Adding Additional Payment Terms**
    
        - Click **Add Payment Split** to create a new row  
        - Enter the percentage and payment condition  
    
        ---
    
        **Important Rules**
    
        - The total of all percentages **must equal 100%**  
        - The system will display a warning if the total is incorrect  
    
        ✅ A green confirmation will appear when the total equals 100%  
    
        ---
    
        ⚠️ Ensure payment descriptions are clear and commercially appropriate, as they will be included in the final quotation document.
        """)

    
    # Display inputs
    for i, term in enumerate(st.session_state.payment_terms):
        col1, col2 = st.columns([1, 3])
    
        with col1:  
            
            if f"percent_{i}" not in st.session_state:
                st.session_state[f"percent_{i}"] = int(term["percent"])
            st.number_input(
                f"% {i+1}",
                min_value=0,
                max_value=100,
                key=f"percent_{i}",
                on_change=update_payment_term,
                args=(i,)
            )
            

        with col2:
            st.session_state.payment_terms[i]["description"] = st.text_input(
                f"Condition {i+1}",
                value=term["description"],
                key=f"desc_{i}",
                on_change=lambda: st.session_state.update({"payment_override": True})
            )
    
    # Add new row
    if st.button("➕ Add Payment Split"):
        st.session_state.payment_terms.append({"percent": 0, "description": ""})
    
    # Calculate total
    total_percent = sum(float(term["percent"]) for term in st.session_state.payment_terms)
    
    if total_percent != 100:
        st.warning(f"⚠️ Total must equal 100% (Currently {total_percent}%)")
    else:
        st.success("✅ Payment terms total = 100%")
                

# ==========================
# 📄 FILE NAMING SECTION
# ==========================
st.markdown("---")
with st.expander("📄 File Naming", expanded=False):
    
    with st.expander("💡 How to fill in", expanded=False):
        st.markdown("""
        **This section determines the name of the generated quotation file.**
    
        Please ensure all fields are selected correctly, as they will be combined to form the final document name.
    
        The filename will follow the format:
        - Project Number  (This should automatically populate if filled in during Template Info section)
        - Document Type  
        - Subject  
        - Unique Identifier  
        - Revision Code and Number  
    
        ✅ Ensure naming is consistent with project conventions.
    
        ⚠️ Incorrect or inconsistent inputs may result in an incorrectly named file.
        """)


    document_type_options = {
        "Cost Sheet": "CST",
        "Quote": "QTE",
        "Calculation": "CLC",
        "Data Collection Form": "DCF",
        "Document": "DOC",
        "Drawing": "DRG",
        "Functional Design Specification": "FDS",
        "Risk Assessment / Method Statement": "RMS",
        "Report": "RPT",
        "Schedule": "SCH",
        "Specification": "SPC"
    }
    
    subject_options = {
        "Audit (EcoConsult Audit for Power)": "ADT",
        "Block": "BLK",
        "Cabling": "CBL",
        "Design": "DES",
        "Data Centre Audit": "DTC",
        "Equipment": "EQP",
        "Earthing": "ETH",
        "Factory Acceptance Test": "FAT",
        "Feasibility": "FSY",
        "General Arrangement or Layout": "GAR",
        "Energy Audit": "NRG",
        "Microgrid Feasibility": "MGF",
        "Microgrid Design": "MGD",
        "Metering Survey": "MTR",
        "Protection": "PRT",
        "Power Quality": "PQT",
        "Pressure Rise Study": "PRS",
        "Power System Study": "PSS"
    }
    
    col1, col2 = st.columns(2)
    
    with col1:
        
            # ✅ Always sync Project Number to Project Name (LIVE)
        st.session_state.project_number = project_name
        
        # ✅ Live sync Project Name → Project Number
        st.session_state.project_number = project_name
        
        project_number = st.text_input(
            "Project Number",
            key="project_number",
            value=st.session_state.project_number
        )
    
    
        document_type_label = st.selectbox("Document Type", list(document_type_options.keys()))
        document_type = document_type_options[document_type_label]
    
        subject_label = st.selectbox("Subject", list(subject_options.keys()))
        subject = subject_options[subject_label]
    
    with col2:
        unique_id = st.selectbox("Unique Identifier", [f"{i:02d}" for i in range(1, 21)])
    
        revision_code_label = st.selectbox(
            "Revision Code",
            ["Contractual (External)", "Preliminary (Internal)"]
        )
        revision_code = "C" if "Contractual" in revision_code_label else "P"
    
        revision_number = st.selectbox("Revision Number", [f"{i:02d}" for i in range(1, 21)])

# ==========================
# TRANSPORT SECTION
# ==========================
st.markdown("---")

if "Transport" in quote_type:
    with st.expander("🚚 Transport Details", expanded=False):

        col1, col2 = st.columns(2)
    
        with col1:
            number_of_transport = st.text_input("Number of Transports")
    
        with col2:
            transport_type = st.text_input(
                "Transport Type",
                placeholder="e.g. HGV and Grey Fleet"
            )
            st.caption("If multiple types, use: HGV and Grey Fleet")
    
else:
    number_of_transport = ""
    transport_type = ""

st.markdown("---")

# ==========================
# ✅ SAFE PLACEHOLDER FUNCTION
# ==========================
   
def replace_placeholders(doc, data):

    def process_paragraph(paragraph):

        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"

            # ✅ Check full text
            full_text = "".join(run.text for run in paragraph.runs)

            if placeholder in full_text:

                # ✅ Replace across runs carefully
                remaining_text = full_text.replace(placeholder, str(value))

                # ✅ Write back character by character across runs
                i = 0
                for run in paragraph.runs:
                    run_len = len(run.text)

                    run.text = remaining_text[i:i+run_len]

                    i += run_len

                # ✅ If any text left, append to last run
                if i < len(remaining_text):
                    paragraph.runs[-1].text += remaining_text[i:]

    # ✅ Process paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # ✅ Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    return doc


def fill_works_table(doc, works_list, fx_rate, currency_symbol):

    target_table = None

    # ✅ Find the table with "Work Description" header
    for table in doc.tables:
        first_row = table.rows[0].cells

        if "Description" in first_row[1].text:
            target_table = table
            break

    if target_table is None:
        print("❌ Pricing table not found")
        return doc

    start_row = 1

    # ✅ Fill rows
    for i, work in enumerate(works_list):

        if start_row + i >= len(target_table.rows) - 1:
            row_cells = target_table.add_row().cells
        else:
            row_cells = target_table.rows[start_row + i].cells

        row_cells[0].text = str(i + 1)
        row_cells[1].text = work["description"]
        price_fx = work["price"] * fx_rate
        row_cells[2].text = f"{currency_symbol}{price_fx:,.2f}"

    # ✅ Total row
    total = sum(work["price"] for work in works_list)
    total_fx = total * fx_rate
    target_table.rows[-1].cells[2].text = f"{currency_symbol}{total_fx:,.2f}"

    return doc



def insert_payment_terms(doc, payment_terms):

    payment_lines = []

    for term in payment_terms:
        if term["percent"] > 0 and term["description"]:
            payment_lines.append(f" {term['percent']}% {term['description']}")

    payment_text = "\n".join(payment_lines)

    # ✅ 1. Normal paragraphs
    for paragraph in doc.paragraphs:
        if "{{PaymentTerms}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{PaymentTerms}}", payment_text)

    # ✅ 2. TABLE CELLS (THIS IS THE MISSING PIECE)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "{{PaymentTerms}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{PaymentTerms}}", payment_text)

    return doc



# ==========================
# GENERATE DOCUMENT
# ==========================
if st.button("📄 Generate Word Document"):
    
    
    if uploaded_file is not None and "office_hours" in locals():
        
        uploaded_file.seek(0)
        file_bytes = uploaded_file.read()
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes))
        ws = wb["PRICING SHEET"]
    
        # ==========================
        # ✅ PRICING SHEET HEADER INFO
        # ==========================
        
        ws["C5"] = sap_description
        ws["C8"] = consultant_name
        ws["J5"] = currency
        ws["B10"] = office_hours
        ws["C10"] = site_hours
        # ==========================
        # ✅ SAP INFO FORM SHEET
        # ==========================
        
        sap_ws = wb["SAP INFO FORM"]
        
        sap_ws["D7"] = customer_contact_name
        sap_ws["D8"] = contact_tel
        sap_ws["D9"] = contact_email
        
        sap_ws["D17"] = bfo_opp_no
        sap_ws["D18"] = material_code_count
        
        sap_ws["D19"] = material_code_1
        sap_ws["D20"] = material_code_2
        sap_ws["D21"] = material_code_3


    if not customer_name or not project_name:
        st.error("Customer Name and Project Name are required.")

    elif not st.session_state.works_list:
        st.error("Please add at least one work item before generating the document.")

    elif not st.session_state.works_list:
        st.error("Please add at least one work item before generating the document.")
    
    elif any(
        not work.get("description") or work.get("price", 0) == 0
        for work in st.session_state.works_list
    ):
        st.error("❌ All work items must have a description and non-zero price before generating the document.")

    else:
        template_path = quote_options[quote_type]
        doc = Document(template_path)

        # ✅ Fix contact name logic
        if not contact_name.strip():
            contact_name_final = "whom it may concern"
        else:
            contact_name_final = contact_name

        # ✅ Build address properly (FIXED INDENT)
        address_lines = []

        if address_line_1:
            address_lines.append(address_line_1)

        if address_line_2:
            address_lines.append(address_line_2)

        if city:
            address_lines.append(city)

        if postcode:
            address_lines.append(postcode)

        full_address = "\n".join(address_lines)

        # ✅ Build data dictionary
        data = {
            "CustomerName": customer_name,
            "NumberOfSites": number_of_sites,
            "SiteName": site_name,
            "NumberOfTransport": number_of_transport,
            "TransportType": transport_type,
            "ProjectName": project_name,
            "NumberOfConsultants": number_of_consultants,

            "FullAddress": full_address,
            "ContactName": contact_name_final,

            "TodaysDate": datetime.now().strftime("%d %B %Y")
        }

        # ✅ Apply transformations
        doc = insert_payment_terms(doc, st.session_state.payment_terms)
        doc = replace_placeholders(doc, data)

        if st.session_state.works_list:
            doc = fill_works_table(doc, st.session_state.works_list, fx_rate, currency_symbol)

        # ✅ Filename
        file_name = f"{project_number}-{document_type}-{subject}-{unique_id}-{revision_code}{revision_number}"

        output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(output.name)

        st.success("✅ Document generated successfully!")

        with open(output.name, "rb") as f:
            st.download_button(
                "⬇ Download Quote",
                f,
                file_name=f"{file_name}.docx"
            )
        
        # ✅ Generate Excel
        
        if uploaded_file is not None:
        
            excel_wb = generate_pricing_excel(uploaded_file)
        
            excel_output = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
            excel_wb.save(excel_output.name)
        
            with open(excel_output.name, "rb") as f:
                st.download_button(
                    "⬇ Download Pricing Sheet",
                    f,
                    file_name=f"{project_number}-PricingSheet.xlsx"
                )

